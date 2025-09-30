# utils/export_utils.py
import io
from docx import Document
from fpdf import FPDF

def create_docx_from_text(text: str) -> io.BytesIO:
    """Create a DOCX document from text"""
    if not text or not text.strip():
        raise ValueError("Cannot create document from empty text")
    
    try:
        doc = Document()
        doc.add_heading("Lecture Notes (Generated)", level=1)
        
        for line in text.splitlines():
            if line.strip():  # Only add non-empty lines
                doc.add_paragraph(line)
        
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    except Exception as e:
        raise Exception(f"DOCX creation failed: {str(e)}")

def create_pdf_from_text(text: str) -> io.BytesIO:
    """Create a PDF document from text"""
    if not text or not text.strip():
        raise ValueError("Cannot create PDF from empty text")
    
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        
        # More aggressive Unicode character replacement
        import re
        # Replace common problematic Unicode characters
        safe_text = text
        safe_text = re.sub(r'["""]', '"', safe_text)  # Replace smart quotes
        safe_text = re.sub(r"[''']", "'", safe_text)  # Replace smart apostrophes  
        safe_text = re.sub(r'[—–]', '-', safe_text)   # Replace em/en dashes
        safe_text = re.sub(r'[…]', '...', safe_text)  # Replace ellipsis
        safe_text = re.sub(r'[•]', '*', safe_text)    # Replace bullet points
        
        # Filter to only ASCII characters as final safety
        safe_text = ''.join(char for char in safe_text if ord(char) < 128)
        
        if not safe_text.strip():
            raise ValueError("No valid text content after character filtering")
        
        for line in safe_text.splitlines():
            if line.strip():  # Only add non-empty lines
                try:
                    pdf.multi_cell(0, 8, line)
                except UnicodeEncodeError:
                    # Skip lines that still cause issues
                    continue
        
        buf = io.BytesIO()
        pdf_output = pdf.output(dest='S')
        buf.write(pdf_output.encode('latin-1'))
        buf.seek(0)
        return buf
    except Exception as e:
        raise Exception(f"PDF creation failed: {str(e)}")
